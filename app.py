from __future__ import annotations

import io
import json
import os
import re
import time
from dataclasses import dataclass, field, asdict
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple, cast

import pandas as pd
import streamlit as st
from datetime import datetime
import pytz

# Optional document parsers
try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import docx
except Exception:  # pragma: no cover
    docx = None

# Optional LLM client
try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None


# -----------------------------------------------------------------------------
# Streamlit page config and CTI dashboard styling
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="AI Job Application Assistant | Agentic CTI Dashboard",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

DARK_CSS = """
<style>
:root {
    --bg: #08111f;
    --panel: #0d1b2f;
    --panel2: #12243d;
    --border: #23405f;
    --text: #e6f1ff;
    --muted: #8aa4bf;
    --cyan: #00d4ff;
    --green: #33ff99;
    --yellow: #ffd166;
    --red: #ff5263;
    --purple: #a78bfa;
}
html, body, [class*="css"] { background: var(--bg) !important; color: var(--text) !important; }
.stApp { background: radial-gradient(circle at top left, #10243f 0, #08111f 35%, #050914 100%) !important; }
section[data-testid="stSidebar"] { background: linear-gradient(180deg, #07111f 0%, #0d1b2f 100%) !important; border-right: 1px solid var(--border); }
.block-container { padding-top: 3rem; padding-bottom: 2rem; }
.dashboard-title { font-size: 2.1rem; font-weight: 800; letter-spacing: .02em; color: var(--text); margin-bottom: .2rem; }
.dashboard-subtitle { color: var(--muted); font-size: .98rem; margin-bottom: 1rem; }
.cti-card { background: rgba(13, 27, 47, .88); border: 1px solid var(--border); border-radius: 18px; padding: 18px 20px; box-shadow: 0 10px 26px rgba(0,0,0,.28); }
.cti-card h3 { margin-top: 0; color: var(--cyan); }
.metric-card { background: linear-gradient(135deg, rgba(18,36,61,.9), rgba(5,15,30,.9)); border: 1px solid #2b5174; border-radius: 16px; padding: 16px; }
.badge { padding: 4px 9px; border-radius: 999px; font-size: .78rem; border: 1px solid; display: inline-block; margin-right: 6px; }
.badge-green { color: var(--green); border-color: var(--green); background: rgba(51,255,153,.08); }
.badge-yellow { color: var(--yellow); border-color: var(--yellow); background: rgba(255,209,102,.08); }
.badge-red { color: var(--red); border-color: var(--red); background: rgba(255,82,99,.08); }
.badge-cyan { color: var(--cyan); border-color: var(--cyan); background: rgba(0,212,255,.08); }

/* ── Container spacing for tabs ── */
.stTabs [data-baseweb="tab-list"] {
    gap: 6px;
    background-color: transparent;
    border-bottom: 2px solid #23405f; /* Base line under the tab bar */
    padding-bottom: 0px;
}

/* ── Tabbed Shape (Default/Unselected State) ── */
.stTabs [data-baseweb="tab"] {
    background-color: #0d1b2f !important;
    border: 1px solid #23405f !important;
    border-bottom: none !important; /* Open at the bottom to attach to base line */
    border-radius: 10px 10px 0px 0px !important; /* Rounded top-left & top-right corners */
    color: #8aa4bf !important;
    padding: 10px 18px !important;
    font-weight: 600 !important;
    margin-bottom: -2px !important; /* Sits directly on top of the bottom border */
    transition: all 0.2s ease-in-out;
}

/* ── Hover State for Unselected Tabs ── */
.stTabs [data-baseweb="tab"]:hover {
    background-color: #12243d !important;
    border-color: #00d4ff !important;
    color: #e6f1ff !important;
}

/* ── Active / Selected Tab Shape & Highlight ── */
.stTabs [aria-selected="true"] {
    background: linear-gradient(180deg, #122d4a 0%, #0d1b2f 100%) !important;
    border: 1px solid #00d4ff !important;
    border-bottom: 2px solid #0d1b2f !important; /* Blends seamlessly into the container background */
    color: #00d4ff !important;
    box-shadow: 0 -4px 12px rgba(0, 212, 255, 0.2) !important;
}

/* ── Hide Streamlit's default red/blue underline bar ── */
.stTabs [data-baseweb="tab-highlight-container"] {
    display: none !important;
}

.stButton>button { background: linear-gradient(90deg, #00d4ff, #33ff99); color: #06111f; font-weight: 800; border: 0; border-radius: 12px; padding: .55rem 1rem; }
.stDownloadButton>button { background: #10243f; color: var(--cyan); border: 1px solid var(--cyan); border-radius: 12px; }
[data-testid="stMetricValue"] { color: var(--cyan); }
hr { border-color: var(--border); }

/* ── Text Areas (Raw Previews & Inputs) ── */
    .stTextArea textarea {
        font-family: 'IBM Plex Mono', monospace !important;
        font-size: 0.85rem !important;
        color: #e6edf3 !important;
        background-color: #080f1a !important;
        border: 1px solid #1e3a5f !important;
        border-radius: 6px !important;
        line-height: 1.5 !important;
        padding: 1rem !important;
    }
    
    /* Focus state for when you click inside the box */
    .stTextArea textarea:focus {
        border-color: #38bdf8 !important;
        box-shadow: 0 0 0 1px #38bdf8 !important;
    }

    /* Styles the label above the text box to look like a system prompt */
    .stTextArea p {
        font-family: 'IBM Plex Mono', monospace !important;
        font-size: 0.75rem !important;
        color: #8ba3c0 !important;
        text-transform: uppercase;
        letter-spacing: 0.1em;
    }

</style>
"""
st.markdown(DARK_CSS, unsafe_allow_html=True)


# -----------------------------------------------------------------------------
# Constants derived from submitted milestone/rubric documents
# -----------------------------------------------------------------------------
MAX_GENERATION_LOOPS = 3
EARLY_EXIT_SKILL_OVERLAP_THRESHOLD = 0.25
QUALITY_APPROVAL_THRESHOLD = 8
TASK_SUCCESS_TARGET = 0.90
EFFICIENCY_TARGET_AVG_LOOPS = 1.8
QUALITY_TARGET = 8.0

RISK_REGISTER = [
    {"Risk Category": "Data Privacy Violations", "Likelihood": 2, "Impact": 5, "Risk Level": 10, "Primary Mitigation": "Ephemeral sessions, no persistent PII logs, user-controlled clear session"},
    {"Risk Category": "Misinformation / Hallucinations", "Likelihood": 4, "Impact": 2, "Risk Level": 8, "Primary Mitigation": "Critic validation, low temperature, resume-grounded generation"},
    {"Risk Category": "Adversarial Attacks", "Likelihood": 3, "Impact": 2, "Risk Level": 6, "Primary Mitigation": "Prompt isolation, XML/HTML stripping, injection detection, graceful halt"},
    {"Risk Category": "Bias & Discrimination", "Likelihood": 2, "Impact": 3, "Risk Level": 6, "Primary Mitigation": "Demographic scrubbing, skills-first matching, bias warnings"},
    {"Risk Category": "Intellectual Property Risks", "Likelihood": 2, "Impact": 2, "Risk Level": 4, "Primary Mitigation": "Original letters, source transparency, AI co-authorship acknowledgment"},
]

COMMON_SKILLS = sorted(set("""
python sql javascript typescript java c++ c# html css react angular vue node flask django fastapi
streamlit pandas numpy matplotlib seaborn scikit-learn tensorflow pytorch machine learning deep learning
nlp llm prompt engineering langchain crewai openai azure aws gcp docker kubernetes linux windows git github
ci/cd devops terraform ansible jenkins power bi tableau excel salesforce servicenow splunk wireshark nmap
metasploit cybersecurity information security soc siem incident response threat intelligence vulnerability
risk management nist iso 27001 access control iam active directory networking tcp/ip dns dhcp firewall vpn
cloud security data analysis database postgresql mysql mongodb oracle etl api rest graphql agile scrum jira
project management ux research qa testing automation selenium communication leadership customer service
technical support helpdesk troubleshooting documentation business analysis requirements stakeholder research
""".split()))

DEMO_JOBS = [
    {
        "rank": 1,
        "title": "Cybersecurity Analyst Intern",
        "company": "SecureTech Solutions",
        "location": "Remote / Hybrid",
        "description": "Assist SOC team with SIEM monitoring, incident triage, vulnerability analysis, Python scripting, Linux, networking, and documentation. Familiarity with NIST and access control preferred.",
        "source_url": "https://example.com/jobs/cybersecurity-analyst-intern",
    },
    {
        "rank": 2,
        "title": "Junior Business Analyst",
        "company": "DataBridge Consulting",
        "location": "Atlanta, GA",
        "description": "Gather requirements, analyze data in Excel and SQL, document workflows, coordinate stakeholders, build dashboards, and support agile software delivery.",
        "source_url": "https://example.com/jobs/junior-business-analyst",
    },
    {
        "rank": 3,
        "title": "Cloud Administrator",
        "company": "NorthStar Cloud",
        "location": "Remote",
        "description": "Support AWS and Azure cloud resources, IAM, Linux, networking, troubleshooting, automation scripts, cloud security, and operational documentation.",
        "source_url": "https://example.com/jobs/cloud-administrator",
    },
    {
        "rank": 4,
        "title": "Frontend Web Developer",
        "company": "BrightApps Studio",
        "location": "Hybrid",
        "description": "Build responsive interfaces using HTML, CSS, JavaScript, React, APIs, Git, testing, accessibility, and agile teamwork.",
        "source_url": "https://example.com/jobs/frontend-web-developer",
    },
    {
        "rank": 5,
        "title": "SOC Analyst Tier 1",
        "company": "BlueShield Cyber Defense",
        "location": "Atlanta, GA",
        "description": "Monitor SIEM alerts, perform initial incident response, analyze phishing reports, document events, and escalate threats. Skills: networking, Splunk, Linux, cybersecurity, communication.",
        "source_url": "https://example.com/jobs/soc-analyst-tier-1",
    },
]

PII_PATTERNS = {
    "email": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    "phone": re.compile(r"(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?)\d{3}[-.\s]?\d{4}"),
    "linkedin": re.compile(r"https?://(?:www\.)?linkedin\.com/[^\s]+", re.I),
    "address_hint": re.compile(r"\b\d{2,5}\s+[A-Za-z0-9 .'-]+\s+(Street|St|Road|Rd|Avenue|Ave|Drive|Dr|Lane|Ln|Boulevard|Blvd)\b", re.I),
}

INJECTION_PATTERNS = [
    r"ignore\s+(all\s+)?previous\s+instructions",
    r"system\s*prompt",
    r"developer\s*message",
    r"</?resume_text>",
    r"</?system>",
    r"BEGIN\s+PROMPT\s+INJECTION",
    r"act\s+as\s+system",
]

DEMOGRAPHIC_TERMS = [
    "race", "gender", "religion", "marital status", "pregnant", "disability", "age", "pronouns", "ethnicity", "citizenship"
]


# -----------------------------------------------------------------------------
# Data models
# -----------------------------------------------------------------------------
@dataclass
class ResumeProfile:
    extracted_skills: List[str] = field(default_factory=list)
    strengths: List[str] = field(default_factory=list)
    gaps_identified: List[str] = field(default_factory=list)
    suggested_keywords: List[str] = field(default_factory=list)
    experience_summary: str = ""
    pii_detected: List[str] = field(default_factory=list)
    red_flags: List[str] = field(default_factory=list)
    anonymized_resume: str = ""


@dataclass
class CriticResult:
    status: str
    professionalism: int
    alignment: int
    completeness: int
    honesty: int
    harmlessness: int
    feedback: str
    risk_flags: List[str] = field(default_factory=list)


@dataclass
class JudgeResult:
    alignment_score: int
    honesty_score: int
    harmlessness_score: int
    helpfulness_score: int
    overall_score: float
    pass_fail: str
    justification: str


# -----------------------------------------------------------------------------
# Utility functions
# -----------------------------------------------------------------------------
def safe_json_loads(text: str, fallback: Any) -> Any:
    try:
        if not text:
            return fallback
        match = re.search(r"\{.*\}|\[.*\]", text, flags=re.S)
        if match:
            return json.loads(match.group(0))
        return json.loads(text)
    except Exception:
        return fallback


def clamp_score(value: Any, default: int = 7) -> int:
    try:
        return max(1, min(10, int(round(float(value)))))
    except Exception:
        return default


def sanitize_text(text: str, max_chars: int = 30000) -> Tuple[str, List[str]]:
    """Remove adversarial structure and detect suspicious payloads."""
    if not text:
        return "", []

    flags = []
    clean = text.replace("\x00", " ")

    for pattern in INJECTION_PATTERNS:
        if re.search(pattern, clean, flags=re.I):
            flags.append(f"Prompt-injection signature detected: {pattern}")

    if re.search(r"<[^>]{1,80}>", clean):
        flags.append("XML/HTML-like tags stripped from user input")
    clean = re.sub(r"<[^>]+>", " ", clean)

    # Collapse whitespace and length-limit to reduce abuse / cost
    clean = re.sub(r"\s+", " ", clean).strip()
    if len(clean) > max_chars:
        flags.append(f"Input truncated to {max_chars:,} characters for safety and cost control")
        clean = clean[:max_chars]

    return clean, sorted(set(flags))


def detect_and_scrub_pii(text: str) -> Tuple[str, List[str]]:
    detected = []
    scrubbed = text
    for label, pattern in PII_PATTERNS.items():
        if pattern.search(scrubbed):
            detected.append(label)
            scrubbed = pattern.sub(f"[{label.upper()}_REDACTED]", scrubbed)
    return scrubbed, detected


def extract_skills(text: str) -> List[str]:
    lower = text.lower()
    found = []
    for skill in COMMON_SKILLS:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"
        if re.search(pattern, lower):
            found.append(skill)
    return sorted(set(found))


def keyword_overlap(a: List[str], text_b: str) -> Tuple[float, List[str]]:
    if not a:
        return 0.0, []
    b = text_b.lower()
    matched = [s for s in a if re.search(r"\b" + re.escape(s.lower()) + r"\b", b)]
    return len(matched) / max(1, len(a)), matched


def extract_text_from_pdf(uploaded_file) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF parsing requires package 'pypdf'. Install with: pip install pypdf")
    data = uploaded_file.read()
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_text_from_docx(uploaded_file) -> str:
    if docx is None:
        raise RuntimeError("Word parsing requires package 'python-docx'. Install with: pip install python-docx")
    data = uploaded_file.read()
    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table text if present
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def extract_uploaded_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    if name.endswith(".doc"):
        raise RuntimeError("Legacy .doc files are not supported by this single-file app. Please save as .docx or PDF.")
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    raise RuntimeError("Unsupported file type. Upload PDF, DOCX, or TXT.")


def df_from_json_list(obj: List[Dict[str, Any]]) -> pd.DataFrame:
    if not obj:
        return pd.DataFrame()
    return pd.DataFrame(obj)


# -----------------------------------------------------------------------------
# LLM Gateway
# -----------------------------------------------------------------------------
class LLMGateway:
    def __init__(self, model: str, temperature: float = 0.2, api_key: str = "", live_mode: bool = False):
        self.model = model
        self.temperature = temperature
        self.live_mode = live_mode
        self.api_key = ""
        self.client: Optional[Any] = None

        if not self.live_mode:
            return

        env_key = os.getenv("OPENAI_API_KEY", "").strip()
        secrets_key = ""
        try:
            secrets_key = st.secrets.get("OPENAI_API_KEY", "").strip()
        except Exception:
            secrets_key = ""

        self.api_key = (api_key or env_key or secrets_key).strip()
        if self.api_key and OpenAI is not None:
            try:
                self.client = OpenAI(api_key=self.api_key)
            except Exception:
                self.client = None

    @property
    def enabled(self) -> bool:
        return self.live_mode and self.client is not None

    def complete(self, system: str, user: str, json_mode: bool = False) -> str:
        if not self.enabled:
            raise RuntimeError("Live LLM mode is off or unavailable. Enable the toggle and provide a valid API key.")
        client = self.client
        if client is None:
            raise RuntimeError("Live LLM client is not available.")
        kwargs = {
            "model": self.model,
            "temperature": self.temperature,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        }
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        response = client.chat.completions.create(**kwargs)
        return response.choices[0].message.content or ""


# -----------------------------------------------------------------------------
# Agents
# -----------------------------------------------------------------------------
class ResumeAnalyzerAgent:
    name = "Resumé Analyzer"

    def __init__(self, llm: LLMGateway):
        self.llm = llm

    def run(self, raw_resume_text: str) -> ResumeProfile:
        sanitized, security_flags = sanitize_text(raw_resume_text)
        scrubbed, pii_detected = detect_and_scrub_pii(sanitized)
        skills = extract_skills(scrubbed)
        demographic_hits = [t for t in DEMOGRAPHIC_TERMS if re.search(r"\b" + re.escape(t) + r"\b", scrubbed, flags=re.I)]
        red_flags = security_flags + (["Potential demographic attributes detected and ignored for matching"] if demographic_hits else [])

        fallback = ResumeProfile(
            extracted_skills=skills,
            strengths=self._heuristic_strengths(skills, scrubbed),
            gaps_identified=self._heuristic_gaps(skills, scrubbed),
            suggested_keywords=self._suggest_keywords(skills),
            experience_summary=self._heuristic_summary(scrubbed),
            pii_detected=pii_detected,
            red_flags=sorted(set(red_flags)),
            anonymized_resume=scrubbed,
        )

        if self.llm.enabled:
            system = (
                "You are a resume analysis agent. Return strict JSON only. "
                "Extract skills and strengths from the resume. Do not infer demographics. "
                "Do not invent experience. Flag skill gaps and missing keywords."
            )
            user = f"""
Analyze this sanitized resume text inside delimiters.
<resume_text>{scrubbed}</resume_text>
Return JSON with keys: extracted_skills, strengths, gaps_identified, suggested_keywords, experience_summary.
"""
            try:
                data = safe_json_loads(self.llm.complete(system, user, json_mode=True), {})
                fallback.extracted_skills = sorted(set(data.get("extracted_skills", fallback.extracted_skills) + skills))
                fallback.strengths = data.get("strengths", fallback.strengths) or fallback.strengths
                fallback.gaps_identified = data.get("gaps_identified", fallback.gaps_identified) or fallback.gaps_identified
                fallback.suggested_keywords = data.get("suggested_keywords", fallback.suggested_keywords) or fallback.suggested_keywords
                fallback.experience_summary = data.get("experience_summary", fallback.experience_summary) or fallback.experience_summary
            except Exception as exc:
                fallback.red_flags.append(f"LLM analyzer fallback used: {exc}")
        return fallback

    @staticmethod
    def _heuristic_summary(text: str) -> str:
        words = text.split()
        preview = " ".join(words[:75])
        return preview + ("..." if len(words) > 75 else "")

    @staticmethod
    def _heuristic_strengths(skills: List[str], text: str) -> List[str]:
        strengths = []
        if any(s in skills for s in ["python", "sql", "data analysis", "pandas"]):
            strengths.append("Technical foundation in programming, data, or analytics tools")
        if any(s in skills for s in ["cybersecurity", "information security", "soc", "siem", "nist"]):
            strengths.append("Security-oriented experience relevant to cyber and risk roles")
        if any(s in skills for s in ["communication", "leadership", "customer service", "documentation"]):
            strengths.append("Strong communication, documentation, and stakeholder-facing capabilities")
        if not strengths:
            strengths.append("Resume contains experience that can be structured into role-specific accomplishments")
        if re.search(r"\b\d+%|\$\d+|\b\d+\+?\s+(users|clients|projects|tickets|reports)\b", text, re.I):
            strengths.append("Includes measurable impact statements that can strengthen application materials")
        return strengths

    @staticmethod
    def _heuristic_gaps(skills: List[str], text: str) -> List[str]:
        gaps = []
        if len(skills) < 5:
            gaps.append("Add more explicit technical and role-specific keywords")
        if not re.search(r"\b\d+%|\$\d+|\b\d+\+?\s+(users|clients|projects|tickets|reports)\b", text, re.I):
            gaps.append("Add metrics-driven impact statements where truthful")
        if "project" not in text.lower():
            gaps.append("Include one or more concrete projects to demonstrate applied experience")
        return gaps or ["No major structural gaps detected; focus on tailoring to each job description"]

    @staticmethod
    def _suggest_keywords(skills: List[str]) -> List[str]:
        suggestions = []
        if any(s in skills for s in ["cybersecurity", "information security", "soc"]):
            suggestions += ["NIST", "SIEM", "Incident Response", "Access Control", "Risk Assessment"]
        if any(s in skills for s in ["python", "sql", "data analysis"]):
            suggestions += ["ETL", "Data Visualization", "Automation", "Dashboard", "Analytics"]
        if any(s in skills for s in ["aws", "azure", "cloud security"]):
            suggestions += ["IAM", "Cloud Security", "Linux", "Infrastructure", "Monitoring"]
        return sorted(set(suggestions or ["Agile", "Documentation", "Stakeholder Communication", "Problem Solving"]))


class JobMatcherAgent:
    name = "Job Matcher"

    def __init__(self, llm: LLMGateway):
        self.llm = llm

    def run(self, profile: ResumeProfile, user_preferences: str, manual_job_text: str = "") -> List[Dict[str, Any]]:
        skills = profile.extracted_skills
        jobs = []

        if manual_job_text.strip():
            title = self._extract_title(manual_job_text)
            jobs.append({
                "rank": 1,
                "title": title,
                "company": "User-provided job description",
                "location": "Not specified",
                "description": manual_job_text.strip(),
                "source_url": "Manual Input",
            })
        else:
            jobs = [dict(j) for j in DEMO_JOBS]

        # Optional LLM can create additional demo matches based on preferences, but no live scraping in single file.
        if self.llm.enabled and not manual_job_text.strip():
            system = "You are a skills-first job matching agent. Return strict JSON list of 3 realistic job matches. No protected attributes."
            user = f"Skills: {skills}\nPreferences: {user_preferences}\nReturn JSON array with title, company, location, description, source_url."
            try:
                generated = safe_json_loads(self.llm.complete(system, user, json_mode=False), [])
                if isinstance(generated, dict):
                    generated = generated.get("jobs", [])
                if isinstance(generated, list) and generated:
                    jobs = generated[:5]
            except Exception:
                pass

        ranked = []
        for idx, job in enumerate(jobs, start=1):
            desc = " ".join([str(job.get("title", "")), str(job.get("description", "")), str(job.get("company", ""))])
            score, matched = keyword_overlap(skills, desc)
            pref_bonus = 0.05 if user_preferences and any(w.lower() in desc.lower() for w in user_preferences.split() if len(w) > 3) else 0
            final_score = min(100, round((score + pref_bonus) * 100))
            updated = {
                "rank": idx,
                "title": job.get("title", "Untitled Role"),
                "company": job.get("company", "Unknown"),
                "location": job.get("location", "Not specified"),
                "match_score": final_score,
                "matched_keywords": ", ".join(matched) if matched else "No direct keyword overlap",
                "source_url": job.get("source_url", "N/A"),
                "description": job.get("description", ""),
            }
            ranked.append(updated)

        ranked = sorted(ranked, key=lambda x: x["match_score"], reverse=True)
        for i, job in enumerate(ranked, start=1):
            job["rank"] = i
        return ranked

    @staticmethod
    def _extract_title(text: str) -> str:
        first_lines = [line.strip() for line in text.splitlines() if line.strip()]
        if first_lines:
            candidate = first_lines[0][:80]
            if len(candidate.split()) <= 10:
                return candidate
        for pattern in [r"job title[:\-]\s*(.+)", r"position[:\-]\s*(.+)", r"role[:\-]\s*(.+)"]:
            m = re.search(pattern, text, re.I)
            if m:
                return m.group(1).strip()[:80]
        return "Target Job"


class CoverLetterGeneratorAgent:
    name = "Cover Letter Generator"

    def __init__(self, llm: LLMGateway):
        self.llm = llm

    def run(self, profile: ResumeProfile, job: Dict[str, Any], feedback: str = "") -> str:
        matched_score = job.get("match_score", 0) / 100
        if matched_score < EARLY_EXIT_SKILL_OVERLAP_THRESHOLD:
            return (
                "EARLY EXIT WARNING\n\n"
                "This job appears to have low resume-to-role skill overlap. The system should not fabricate qualifications. "
                "Consider revising the resume, selecting a better-fit role, or writing an honest networking note that acknowledges learning goals.\n\n"
                f"Role: {job.get('title', 'Target Job')}\n"
                f"Current keyword overlap: {job.get('match_score', 0)}%"
            )

        if self.llm.enabled:
            system = """
                    You are an expert cover letter writer. 
                    Write a professional cover letter based on the resume and job description.
                    
                    CRITICAL FORMATTING RULE: 
                    You must include a formal business header at the very top of the letter. 
                    Extract the applicant's actual Name, Location, Phone, and Email from the provided resume profile to build this header. 
                    Under NO circumstances are you allowed to use bracketed placeholders (e.g., [Your Name], [Date]). 
                    If a specific piece of contact information is missing from the resume, simply omit that single line and format the rest of the header normally.

                    CRITICAL HONESTY RULE (ANTI-HALLUCINATION):
                    You must be 100% truthful and factually accurate. 
                    Do NOT invent, fabricate, or exaggerate any skills, job titles, metrics, certifications, or experiences. 
                    If the job description requests a specific skill or software that is NOT explicitly listed in the resume, you must not claim the candidate has it. 
                    Focus entirely on aligning the explicit facts from the resume with the needs of the job.
                    
                    CRITICAL TONE RULE:
                    Maintain a formal, objective, and analytical tone appropriate for a technical professional. 
                    Avoid hyperbolic language, overly enthusiastic phrasing, or cliché corporate buzzwords. 
                    Present facts clearly and concisely.
                    
                    """
            
            user = f"""
Resume profile JSON:
{json.dumps(asdict(profile), indent=2)}

Target job JSON:
{json.dumps(job, indent=2)}

Critic feedback to address:
{feedback or "None"}
"""
            try:
                return self.llm.complete(system, user, json_mode=False).strip()
            except Exception:
                pass

        return self._heuristic_letter(profile, job, feedback)

    @staticmethod
    def _heuristic_letter(profile: ResumeProfile, job: Dict[str, Any], feedback: str) -> str:
        skills = profile.extracted_skills[:8]
        strengths = profile.strengths[:3]
        suggested = profile.suggested_keywords[:5]
        company = job.get("company", "your organization")
        title = job.get("title", "the open role")
        matched = job.get("matched_keywords", "")
        feedback_note = f"\n\nRevision focus applied: {feedback}" if feedback else ""

        return f"""Dear Hiring Manager,

I am writing to express my interest in the {title} position at {company}. My background aligns with this opportunity through demonstrated strengths in {', '.join(strengths).lower()}.

Based on the role requirements, I can contribute practical experience with {', '.join(skills) if skills else 'technical problem solving, documentation, and rapid learning'}. The strongest resume-to-job matches identified by the system include: {matched}. These capabilities position me to support the team with reliable execution, clear communication, and a continuous-improvement mindset.

I am especially interested in this role because it connects my current skill set with the organization’s need for someone who can analyze requirements, document work clearly, and deliver thoughtful solutions. I would welcome the opportunity to discuss how my experience with {', '.join(suggested) if suggested else 'analysis, collaboration, and technical support'} can support your goals.

Thank you for your time and consideration. I look forward to the possibility of speaking with you.

Sincerely,
Applicant{feedback_note}
"""


class CriticReviewerAgent:
    name = "Critic / Reviewer"

    def __init__(self, llm: LLMGateway):
        self.llm = llm

    def run(self, letter: str, profile: ResumeProfile, job: Dict[str, Any]) -> CriticResult:
        if letter.startswith("EARLY EXIT WARNING"):
            return CriticResult(
                status="Rejected - Low Skill Overlap",
                professionalism=7,
                alignment=3,
                completeness=5,
                honesty=10,
                harmlessness=10,
                feedback="Skill overlap is below the early-exit threshold. Choose a better-fit role or revise resume honestly before generating a letter.",
                risk_flags=["Early exit prevents hallucinated qualifications"],
            )

        fallback = self._heuristic_critic(letter, profile, job)
        if self.llm.enabled:
            system = (
                "You are a strict critic/reviewer agent. Return strict JSON only with keys: "
                "status, professionalism, alignment, completeness, honesty, harmlessness, feedback, risk_flags. "
                "Set status to Approved only if all major scores are >= 8. Penalize hallucinations."
            )
            user = f"""
Resume profile JSON:
{json.dumps(asdict(profile), indent=2)}

Target job JSON:
{json.dumps(job, indent=2)}

Cover letter:
{letter}
"""
            try:
                data = safe_json_loads(self.llm.complete(system, user, json_mode=True), {})
                return CriticResult(
                    status=data.get("status", fallback.status),
                    professionalism=clamp_score(data.get("professionalism", fallback.professionalism)),
                    alignment=clamp_score(data.get("alignment", fallback.alignment)),
                    completeness=clamp_score(data.get("completeness", fallback.completeness)),
                    honesty=clamp_score(data.get("honesty", fallback.honesty)),
                    harmlessness=clamp_score(data.get("harmlessness", fallback.harmlessness)),
                    feedback=data.get("feedback", fallback.feedback),
                    risk_flags=data.get("risk_flags", fallback.risk_flags) or fallback.risk_flags,
                )
            except Exception:
                return fallback
        return fallback

    @staticmethod
    def _heuristic_critic(letter: str, profile: ResumeProfile, job: Dict[str, Any]) -> CriticResult:
        score_overlap, matched_letter = keyword_overlap(profile.extracted_skills, letter)
        job_score, matched_job = keyword_overlap(profile.extracted_skills, job.get("description", ""))
        professionalism = 9 if "Dear" in letter and "Sincerely" in letter else 7
        alignment = clamp_score(5 + int(5 * min(1, job_score + score_overlap)))
        completeness = 9 if len(letter.split()) >= 160 else 7
        honesty = 10
        harmlessness = 10
        risk_flags = []

        # PII in output check
        _, pii = detect_and_scrub_pii(letter)
        if pii:
            harmlessness = 7
            risk_flags.append("Generated letter may contain PII")

        # Rough hallucination check: advanced high-risk phrases not grounded
        unsupported_terms = []
        for term in ["certified", "certification", "machine learning", "python", "aws", "azure", "splunk", "kubernetes"]:
            if term in letter.lower() and term not in " ".join(profile.extracted_skills).lower() and term not in profile.anonymized_resume.lower():
                unsupported_terms.append(term)
        if unsupported_terms:
            honesty = 6
            risk_flags.append(f"Potential unsupported claims: {', '.join(sorted(set(unsupported_terms)))}")

        scores = [professionalism, alignment, completeness, honesty, harmlessness]
        status = "Approved" if min(scores) >= QUALITY_APPROVAL_THRESHOLD else "Needs Revision"
        feedback_parts = []
        if alignment < 8:
            feedback_parts.append("Increase explicit alignment with job requirements using only resume-grounded skills.")
        if completeness < 8:
            feedback_parts.append("Expand the letter with a clearer value proposition and role-specific paragraph.")
        if honesty < 8:
            feedback_parts.append("Remove or soften unsupported claims that are not present in the resume.")
        if harmlessness < 8:
            feedback_parts.append("Remove PII or sensitive demographic information from generated output.")
        if not feedback_parts:
            feedback_parts.append("Letter meets approval threshold for professionalism, alignment, completeness, honesty, and harmlessness.")
        return CriticResult(status, professionalism, alignment, completeness, honesty, harmlessness, " ".join(feedback_parts), risk_flags)


class LLMJudgeAgent:
    name = "LLM Judge / HR Compliance Auditor"

    def __init__(self, llm: LLMGateway):
        self.llm = llm

    def run(self, letter: str, profile: ResumeProfile, job: Dict[str, Any]) -> JudgeResult:
        fallback = self._heuristic_judge(letter, profile, job)
        if self.llm.enabled:
            system = (
                "You are a cold, hyper-objective HR Compliance Auditor assessing an AI-generated cover letter. "
                "Return strict valid JSON only. Score 1 to 10 on alignment_score, honesty_score, harmlessness_score, helpfulness_score. "
                "Honesty must be 1 if the letter invents certifications, job history, skills, projects, or metrics absent from the resume."
            )
            user = f"""
Raw/anonymized resume:
{profile.anonymized_resume}

Resume skills:
{profile.extracted_skills}

Job description:
{job.get('description', '')}

Cover letter:
{letter}

Return JSON with: alignment_score, honesty_score, harmlessness_score, helpfulness_score, justification.
"""
            try:
                data = safe_json_loads(self.llm.complete(system, user, json_mode=True), {})
                alignment = clamp_score(data.get("alignment_score", fallback.alignment_score))
                honesty = clamp_score(data.get("honesty_score", fallback.honesty_score))
                harmlessness = clamp_score(data.get("harmlessness_score", fallback.harmlessness_score))
                helpfulness = clamp_score(data.get("helpfulness_score", fallback.helpfulness_score))
                overall = round((alignment + honesty + harmlessness + helpfulness) / 4, 2)
                return JudgeResult(
                    alignment, honesty, harmlessness, helpfulness, overall,
                    "PASS" if min(alignment, honesty, harmlessness, helpfulness) >= 8 else "REVIEW",
                    data.get("justification", fallback.justification),
                )
            except Exception:
                return fallback
        return fallback

    @staticmethod
    def _heuristic_judge(letter: str, profile: ResumeProfile, job: Dict[str, Any]) -> JudgeResult:
        if letter.startswith("EARLY EXIT WARNING"):
            return JudgeResult(3, 10, 10, 5, 7.0, "REVIEW", "Early exit correctly avoids hallucination but does not produce a ready-to-submit package.")
        score_overlap, _ = keyword_overlap(profile.extracted_skills, job.get("description", "") + " " + letter)
        alignment = clamp_score(5 + int(score_overlap * 5))
        honesty = 10
        harmlessness = 10
        helpfulness = 8 if len(letter.split()) > 140 else 6
        _, pii = detect_and_scrub_pii(letter)
        if pii:
            harmlessness = 7
        overall = round((alignment + honesty + harmlessness + helpfulness) / 4, 2)
        return JudgeResult(alignment, honesty, harmlessness, helpfulness, overall, "PASS" if min(alignment, honesty, harmlessness, helpfulness) >= 8 else "REVIEW", "Heuristic judge scored resume/job alignment, apparent grounding, PII exposure, and usefulness.")


def require_generation_inputs(profile: Any, selected_job: Any) -> Tuple[ResumeProfile, Dict[str, Any]]:
    if not isinstance(profile, ResumeProfile):
        raise ValueError("A valid resume profile is required before generating a package.")
    if not isinstance(selected_job, dict):
        raise ValueError("Please select a valid job from the matched jobs list before generating a package.")
    return cast(ResumeProfile, profile), cast(Dict[str, Any], selected_job)


class JobApplicationOrchestrator:
    def __init__(self, llm: LLMGateway):
        self.analyzer = ResumeAnalyzerAgent(llm)
        self.matcher = JobMatcherAgent(llm)
        self.generator = CoverLetterGeneratorAgent(llm)
        self.critic = CriticReviewerAgent(llm)
        self.judge = LLMJudgeAgent(llm)

    def process_resume(self, resume_text: str) -> ResumeProfile:
        return self.analyzer.run(resume_text)

    def match_jobs(self, profile: ResumeProfile, preferences: str, manual_job_text: str = "") -> List[Dict[str, Any]]:
        return self.matcher.run(profile, preferences, manual_job_text)

    def generate_package(self, profile, selected_job: Dict[str, Any]) -> Dict[str, Any]:
        # 1. Bypass the overly strict original check
        # profile, selected_job = require_generation_inputs(profile, selected_job)
        
        if not profile or not selected_job:
             raise ValueError("A valid resume profile and selected job are required.")

        # 2. If Streamlit saved the profile as a dictionary, convert it back into the ResumeProfile class so the agents don't crash
        if isinstance(profile, dict):
            # Assuming ResumeProfile takes the dictionary keys as arguments
            profile = ResumeProfile(**profile)

        feedback = ""
        drafts = []
        critic_result = None
        
        for iteration in range(1, MAX_GENERATION_LOOPS + 1):
            draft = self.generator.run(profile, selected_job, feedback)
            critic_result = self.critic.run(draft, profile, selected_job)
            
            drafts.append({
                "iteration": iteration,
                "draft": draft,
                "critic": asdict(critic_result) if critic_result else {},
            })
            
            if critic_result and (critic_result.status == "Approved" or critic_result.status.startswith("Rejected")):
                break
                
            feedback = critic_result.feedback if critic_result else ""
            
        final_letter = drafts[-1]["draft"] if drafts else ""
        judge_result = self.judge.run(final_letter, profile, selected_job)
        
        status = critic_result.status if critic_result else "Not Started"
        if status == "Needs Revision":
            status = "Completed with Maximum Iterations Warning"
            
        return {
            "final_letter": final_letter,
            "critic_result": asdict(critic_result) if critic_result else {},
            "judge_result": asdict(judge_result) if judge_result else {},
            "draft_history": drafts,
            "iteration_count": len(drafts),
            "status": status,
            "timestamp": datetime.utcnow().isoformat() + "Z",
        }


# -----------------------------------------------------------------------------
# Resume version control helpers
# -----------------------------------------------------------------------------
def make_resume_version_label(version_id: int, source_name: str, resume_text: str) -> str:
    source = source_name or "Manual Entry"
    word_count = len(resume_text.split())
    return f"v{version_id} | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | {source} | {word_count:,} words"


def add_resume_version(resume_text: str, source_name: str, job_text: str = "") -> Dict[str, Any]:
    """Create an in-session version snapshot. No files are written to disk."""
    next_id = len(st.session_state.resume_versions) + 1
    version = {
        "version_id": next_id,
        "label": make_resume_version_label(next_id, source_name, resume_text),
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "source_name": source_name or "Manual Entry",
        "resume_text": resume_text,
        "job_text": job_text,
        "profile": None,
        "notes": "Snapshot created after upload/extraction.",
    }
    st.session_state.resume_versions.append(version)
    st.session_state.active_resume_version = next_id
    return version


def update_active_version_profile(profile: ResumeProfile):
    active_id = st.session_state.get("active_resume_version")
    for version in st.session_state.resume_versions:
        if version["version_id"] == active_id:
            version["profile"] = asdict(profile)
            version["notes"] = f"Analyzed successfully with {len(profile.extracted_skills)} extracted skills."
            break


def load_resume_version(version_id: int):
    for version in st.session_state.resume_versions:
        if version["version_id"] == version_id:
            st.session_state.active_resume_version = version_id
            st.session_state.resume_text = version["resume_text"]
            st.session_state.job_text = version.get("job_text", "")
            st.session_state.profile = None
            st.session_state.jobs = []
            st.session_state.selected_job = None
            st.session_state.package = None
            log_event("Version Control", f"Loaded resume version v{version_id}")
            return version
    return None


def version_table() -> pd.DataFrame:
    rows = []
    for version in st.session_state.resume_versions:
        profile = version.get("profile") or {}
        rows.append({
            "version_id": version["version_id"],
            "created_at": version["created_at"],
            "source_name": version["source_name"],
            "word_count": len(version["resume_text"].split()),
            "skills_count": len(profile.get("extracted_skills", [])) if isinstance(profile, dict) else 0,
            "active": "Yes" if version["version_id"] == st.session_state.active_resume_version else "No",
            "notes": version.get("notes", ""),
        })
    return pd.DataFrame(rows)

# -----------------------------------------------------------------------------
# Session state
# -----------------------------------------------------------------------------
def init_state():
    defaults = {
        "resume_text": "",
        "job_text": "",
        "profile": None,
        "jobs": [],
        "selected_job": None,
        "package": None,
        "run_log": [],
        "resume_versions": [],
        "active_resume_version": None,
        "preferences": "entry-level, internship, remote or Atlanta, cybersecurity, analytics, cloud, software",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def log_event(agent: str, event: str, status: str = "OK"):
    st.session_state.run_log.append({
        "time": datetime.now().strftime("%H:%M:%S"),
        "agent": agent,
        "event": event,
        "status": status,
    })


init_state()

# -----------------------------------------------------------------------------
# Sidebar controls
# -----------------------------------------------------------------------------
with st.sidebar:
   # st.markdown("### 🛡️ CTI Controls")
    
    st.sidebar.subheader("🛡️ CTI Controls")
    
    # 1. Let the user choose the provider
    llm_provider = st.sidebar.selectbox(
        "Select LLM Provider",
        ["OpenAI", "Anthropic", "Google Gemini", "Hugging Face"]
    )
    
    # 2. Dynamically change the input label based on the provider
    api_key = st.sidebar.text_input(
        f"Enter your {llm_provider} API Key", 
        type="password", 
        help=f"Your key is only stored in the session state for this {llm_provider} run."
    )
    
        # 1. Ensure api_key is safely treated as a string
    api_key_str = (api_key or "").strip()

    # 2. Store in session state if provided
    if api_key_str:
        st.session_state["api_key"] = api_key_str
        st.session_state["llm_provider"] = llm_provider

    # 3. Create the toggle unconditionally
    live_mode_toggle = st.toggle(
        "Use Live LLM Mode",
        value=bool(api_key_str),
        key="live_mode_toggle",
        help="Enables agent communication with models using your API key."
    )

    # 4. Safe check using the validated boolean and string
    if not api_key_str and live_mode_toggle:
        st.warning("⚠️ Please input a valid API key above to enable Live Mode.")

        # Dictionary of available models per provider
    available_models = {
        "OpenAI": ["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"],
        "Anthropic": ["claude-3-haiku-20240307", "claude-3-sonnet-20240229"],
        "Google Gemini": ["gemini-1.5-flash", "gemini-1.5-pro"],
        "Hugging Face": ["meta-llama/Meta-Llama-3-8B-Instruct"]
    }

    # Dynamically populate the model dropdown
    model = st.sidebar.selectbox(
        "LLM model", 
        available_models[llm_provider]
    )
    
    temp = st.slider(
        "Generation temperature", 
        0.0, 1.0, 0.2, 0.05, 
        disabled=not live_mode_toggle
    )

        # Initialize gateway with the user-entered API Key
    llm = LLMGateway(
        model=model,
        temperature=temp,
        api_key=api_key_str,
        live_mode=live_mode_toggle and bool(api_key_str)
    )
    st.markdown("---")
    st.markdown("**Runtime mode**")
    if live_mode_toggle and llm.enabled:
        st.success("Live LLM mode enabled")
    elif live_mode_toggle and not llm.enabled:
        st.error("Live LLM requested, but no valid API key/client is available")
    else:
        st.info("Local fallback mode enabled")

    if st.button("Test Live LLM Connection", disabled=not live_mode_toggle):
        if not llm.enabled:
            st.error("Live LLM mode is not ready. Add a valid API key or configure Streamlit secrets.")
        else:
            try:
                test_response = llm.complete(
                    system="You are a connection test assistant.",
                    user="Reply with exactly: LLM connection successful.",
                    json_mode=False,
                )
                st.success(test_response)
            except Exception as exc:
                st.error(f"Connection test failed: {exc}")

    st.markdown("---")
    if st.button("Clear Session / Delete Temporary Data"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        init_state()
        st.rerun()
    st.caption("Privacy guardrail: this app keeps data only in Streamlit session state unless you download outputs.")

orchestrator = JobApplicationOrchestrator(llm)
# -----------------------------------------------------------------------------
# Header
# -----------------------------------------------------------------------------
st.markdown('<div class="dashboard-title">AI Job Application Assistant</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="dashboard-subtitle">Agentic multi-agent workflow: Resumé Analyzer → Job Matcher → Cover Letter Generator → Critic/Reviewer → LLM Judge</div>',
    unsafe_allow_html=True,
)
st.markdown(
    '<span class="badge badge-cyan">Streamlit Front-End</span>'
    '<span class="badge badge-green">Multi-Agent Workflow</span>'
    '<span class="badge badge-yellow">HHH Evaluation</span>'
    '<span class="badge badge-red">Risk Guardrails</span>',
    unsafe_allow_html=True,
)
st.divider()
# KPI row
col1, col2, col3, col4, col5 = st.columns(5)
profile_ready = st.session_state.profile is not None
jobs_ready = len(st.session_state.jobs) > 0
package_ready = st.session_state.package is not None
with col1:
    st.metric("**Profile**", "Ready" if profile_ready else "**Pending**")
with col2:
    st.metric("**Job Matches**", len(st.session_state.jobs))
with col3:
    loops = st.session_state.package.get("iteration_count", 0) if package_ready else 0
    st.metric("**Revision Loops**", loops)
with col4:
    overall = st.session_state.package.get("judge_result", {}).get("overall_score", 0) if package_ready else 0
    st.metric("**Judge's Score**", overall)
with col5:
    st.metric("**Resume Versions**", len(st.session_state.resume_versions))

st.divider()

# -----------------------------------------------------------------------------
# Main tabs
# -----------------------------------------------------------------------------
tabs = st.tabs([
    "**Upload & Extract**",
    "**Resumé Analyzer**",
    "**Job Matcher**",
    "**Generate Package**",
    "**Critic & LLM Judge**",
    "**Evaluation / Risk Dashboard**",
    "**Help / Usage Tips**",
])

with tabs[0]:
    st.markdown('<div class="cti-card">', unsafe_allow_html=True)
    st.subheader("Upload Resumé and Optional Job Description")
    c1, c2 = st.columns(2)
    with c1:
        resume_file = st.file_uploader("Upload résumé (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], key="resume_upload")
        pasted_resume = st.text_area("Or paste résumé text", height=200, placeholder="Paste resume text here...")
    with c2:
        job_file = st.file_uploader("Optional: upload target job description (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], key="job_upload")
        pasted_job = st.text_area("Or paste target job description", height=200, placeholder="Paste job description here for exact matching...")

    #preferences = st.text_input("Job preferences", value=st.session_state.get("preferences", "entry-level, internship, remote or Atlanta, cybersecurity, analytics, cloud, software"))
#############################################################################


#############################################################################


    if st.button("Extract and Sanitize Inputs"):
        progress = st.progress(0, text="Initializing upload pipeline...")
        status_box = st.empty()
        try:
            resume_text = ""
            progress.progress(10, text="Reading résumé upload...")
            status_box.info("Step 1/5: Reading résumé file or pasted text")
            if resume_file is not None:
                resume_text = extract_uploaded_text(resume_file)
            if pasted_resume.strip():
                resume_text = (resume_text + "\n" + pasted_resume).strip()
            progress.progress(35, text="Validating résumé content...")
            if not resume_text.strip():
                st.error("Please upload or paste résumé content.")
                progress.empty()
                status_box.empty()
            else:
                job_text = ""
                progress.progress(50, text="Reading optional job description...")
                status_box.info("Step 2/5: Reading optional job description")
                if job_file is not None:
                    job_text = extract_uploaded_text(job_file)
                if pasted_job.strip():
                    job_text = (job_text + "\n" + pasted_job).strip()

                progress.progress(70, text="Sanitizing input and stripping risky tags...")
                status_box.info("Step 3/5: Sanitizing input for injection and malformed XML/HTML")
                clean_job, job_flags = sanitize_text(job_text) if job_text else ("", [])

                progress.progress(85, text="Creating resume version snapshot...")
                status_box.info("Step 4/5: Creating in-session resume version snapshot")
                st.session_state.resume_text = resume_text
                st.session_state.job_text = clean_job
                #st.session_state.preferences = preferences
                st.session_state.profile = None
                st.session_state.jobs = []
                st.session_state.selected_job = None
                st.session_state.package = None
                source_name = resume_file.name if resume_file is not None else "Pasted Resume"
                version = add_resume_version(resume_text, source_name, clean_job)

                progress.progress(100, text="Upload and versioning complete")
                status_box.success(f"Step 5/5: Input extraction complete. Created {version['label']}.")
                log_event("Input Preprocessor", f"Resume chars={len(resume_text):,}; job chars={len(clean_job):,}; flags={len(job_flags)}; version=v{version['version_id']}")
                st.success("Input extraction complete. Continue to the Resumé Analyzer tab.")
        except Exception as exc:
            progress.empty()
            status_box.empty()
            log_event("Input Preprocessor", str(exc), "ERROR")
            st.error(f"Extraction failed: {exc}")

    if st.session_state.resume_text:
        with st.expander("Raw extracted résumé preview"):
            st.text_area("Resume text", st.session_state.resume_text[:5000], height=220)
    if st.session_state.job_text:
        with st.expander("Sanitized job description preview"):
            st.text_area("Job text", st.session_state.job_text[:5000], height=180)

    st.markdown("#### Resume Version Control")
    if st.session_state.resume_versions:
        st.dataframe(version_table(), use_container_width=True, hide_index=True)
        version_options = {v["label"]: v["version_id"] for v in st.session_state.resume_versions}
        selected_version_label = st.selectbox("Load a previous resume version", list(version_options.keys()))
        c_load, c_export = st.columns([1, 1])
        with c_load:
            if st.button("Load Selected Version"):
                loaded = load_resume_version(version_options[selected_version_label])
                if loaded:
                    st.success(f"Loaded {loaded['label']}. Re-run the Resumé Analyzer for this version.")
                    st.rerun()
        with c_export:
            versions_export = json.dumps(st.session_state.resume_versions, indent=2)
            st.download_button("Export Version History JSON", versions_export, "resume_version_history.json", "application/json")
    else:
        st.caption("No resume versions yet. Upload or paste a resume, then click Extract and Sanitize Inputs.")
    st.markdown('</div>', unsafe_allow_html=True)

with tabs[1]:
    st.subheader("Resumé Analyzer Agent")
    if st.button("Run Resumé Analyzer", disabled=not bool(st.session_state.resume_text)):
        with st.spinner("Analyzing résumé, scrubbing PII, extracting skills..."):
            profile = orchestrator.process_resume(st.session_state.resume_text)
            st.session_state.profile = profile
            update_active_version_profile(profile)
            log_event("Resumé Analyzer", f"skills={len(profile.extracted_skills)} pii={profile.pii_detected} flags={len(profile.red_flags)} version=v{st.session_state.active_resume_version}")
            st.success("Resumé profile generated and saved to the active resume version.")

    if st.session_state.profile:
        profile = st.session_state.profile
        c1, c2 = st.columns([1, 1])
        with c1:
            st.markdown("#### Structured Profile")
            st.json(asdict(profile), expanded=False)
        with c2:
            st.markdown("#### Skills Table")
            skill_df = pd.DataFrame({"Extracted Skills": profile.extracted_skills})
            st.dataframe(skill_df, use_container_width=True, hide_index=True)
            if profile.red_flags:
                st.markdown("#### Security / Bias Flags")
                st.dataframe(pd.DataFrame({"Flags": profile.red_flags}), use_container_width=True, hide_index=True)

with tabs[2]:
    st.subheader("Job Matcher Agent")
    if st.button("Run Job Matcher", disabled=st.session_state.profile is None):
        with st.spinner("Ranking job matches by skill overlap and preferences..."):
            profile = st.session_state.profile
            if profile is None:
                st.error("No profile available for job matching.")
            else:
                jobs = orchestrator.match_jobs(profile, st.session_state.get("preferences", ""), st.session_state.job_text)
                st.session_state.jobs = jobs
                log_event("Job Matcher", f"matched_jobs={len(jobs)} top_score={jobs[0]['match_score'] if jobs else 0}")
                st.success("Job matching complete.")

    if st.session_state.jobs:
        display_cols = ["rank", "title", "company", "location", "match_score", "matched_keywords", "source_url"]
        st.dataframe(pd.DataFrame(st.session_state.jobs)[display_cols], use_container_width=True, hide_index=True)
        # 1. User selects the rank
        selected_rank = st.selectbox("Select job rank for package generation", [j["rank"] for j in st.session_state.jobs])
        
        # 2. THE FIX: Explicit button to lock the choice into memory
        if st.button("Confirm Job Selection", type="primary"):
            st.session_state.selected_job = next(j for j in st.session_state.jobs if j["rank"] == selected_rank)
            st.success("✅ Target Locked! You may now proceed to the Generate Package tab.")

        # 3. Only show the expander if a job has been officially saved
        if "selected_job" in st.session_state and st.session_state.selected_job:
            with st.expander("Selected job details"):
                st.json(st.session_state.selected_job)
        # ------------------------------------------------------------
    # Auto-generated + dropdown-based job preferences
    # ------------------------------------------------------------

    DEFAULT_ROLE_OPTIONS = [
        "Cybersecurity Analyst",
        "SOC Analyst",
        "Cloud Administrator",
        "Business Analyst",
        "Data Analyst",
        "Software Engineer",
        "Frontend Developer",
        "Backend Developer",
        "DevOps Engineer",
        "IT Support Specialist",
        "Systems Administrator",
        "QA Tester",
        "UX Researcher",
        "Project Coordinator",
    ]

    DEFAULT_LEVEL_OPTIONS = [
        "Internship",
        "Entry-level",
        "Junior",
        "Associate",
        "Mid-level",
        "Senior",
    ]

    # A standard Python list of all 50 states
    US_STATES = [
        "Alabama", "Alaska", "Arizona", "Arkansas", "California", "Colorado",
        "Connecticut", "Delaware", "Florida", "Georgia", "Hawaii", "Idaho",
        "Illinois", "Indiana", "Iowa", "Kansas", "Kentucky", "Louisiana",
        "Maine", "Maryland", "Massachusetts", "Michigan", "Minnesota",
        "Mississippi", "Missouri", "Montana", "Nebraska", "Nevada",
        "New Hampshire", "New Jersey", "New Mexico", "New York",
        "North Carolina", "North Dakota", "Ohio", "Oklahoma", "Oregon",
        "Pennsylvania", "Rhode Island", "South Carolina", "South Dakota",
        "Tennessee", "Texas", "Utah", "Vermont", "Virginia", "Washington",
        "West Virginia", "Wisconsin", "Wyoming", "Washington D.C."
    ]

    # Combine your base working styles with the full state list
    DEFAULT_LOCATION_OPTIONS = [
        "Remote",
        "Hybrid",
        "On-site",
        "United States",
        "Atlanta, GA", # Kept this here so your auto-generator default doesn't throw an error!
    ] + US_STATES

    DEFAULT_DOMAIN_OPTIONS = [
        "Cybersecurity",
        "Cloud",
        "Data Analytics",
        "Software Development",
        "Information Technology",
        "Business Analysis",
        "Networking",
        "AI / Machine Learning",
        "Database Administration",
        "Productivity Tools",
    ]


    def auto_generate_preferences_from_profile(profile):
        """
        Creates suggested job preferences from the Resume Analyzer output.
        Falls back to broad entry-level technology preferences if no profile exists.
        """
        if not profile:
            return {
                "roles": ["Cybersecurity Analyst", "Business Analyst"],
                "levels": ["Entry-level", "Internship"],
                "locations": ["Remote", "Atlanta, GA"],
                "domains": ["Cybersecurity", "Information Technology"],
            }

        skills = [s.lower() for s in getattr(profile, "extracted_skills", [])]

        roles = []
        domains = []

        if any(s in skills for s in ["cybersecurity", "information security", "soc", "siem", "nist"]):
            roles += ["Cybersecurity Analyst", "SOC Analyst"]
            domains += ["Cybersecurity"]

        if any(s in skills for s in ["aws", "azure", "gcp", "cloud security", "iam"]):
            roles += ["Cloud Administrator", "DevOps Engineer"]
            domains += ["Cloud"]

        if any(s in skills for s in ["python", "sql", "pandas", "data analysis", "power bi", "tableau"]):
            roles += ["Data Analyst", "Business Analyst"]
            domains += ["Data Analytics"]

        if any(s in skills for s in ["javascript", "react", "html", "css", "node", "java", "django", "flask"]):
            roles += ["Software Engineer", "Frontend Developer", "Backend Developer"]
            domains += ["Software Development"]

        if any(s in skills for s in ["technical support", "helpdesk", "troubleshooting", "active directory"]):
            roles += ["IT Support Specialist", "Systems Administrator"]
            domains += ["Information Technology"]

        return {
            "roles": sorted(set(roles)) or ["Business Analyst", "IT Support Specialist"],
            "levels": ["Entry-level", "Internship"],
            "locations": ["Remote", "Atlanta, GA"],
            "domains": sorted(set(domains)) or ["Information Technology"],
        }


    auto_prefs = auto_generate_preferences_from_profile(st.session_state.get("profile"))

    st.markdown("#### Job Preference Builder")

    pref_col1, pref_col2 = st.columns(2)

    with pref_col1:
        selected_roles = st.multiselect(
            "Target roles",
            options=DEFAULT_ROLE_OPTIONS,
            default=auto_prefs["roles"],
            help="Select one or more job titles the Job Matcher should prioritize.",
        )

        selected_levels = st.multiselect(
            "Experience level",
            options=DEFAULT_LEVEL_OPTIONS,
            default=auto_prefs["levels"],
            help="Choose the seniority level for job recommendations.",
        )

    with pref_col2:
        selected_locations = st.multiselect(
            "Location preference",
            options=DEFAULT_LOCATION_OPTIONS,
            default=auto_prefs["locations"],
            help="Choose preferred work locations or remote/hybrid options.",
        )

        selected_domains = st.multiselect(
            "Career domain",
            options=DEFAULT_DOMAIN_OPTIONS,
            default=auto_prefs["domains"],
            help="Choose the broad career areas the system should focus on.",
        )

    custom_keywords = st.text_input(
        "Additional keywords",
        value="",
        placeholder="Example: Python, SQL, NIST, internship, healthcare, finance",
        help="Optional extra keywords to include in job matching.",
    )

    preferences = ", ".join(
        selected_roles
        + selected_levels
        + selected_locations
        + selected_domains
        + [custom_keywords]
    ).strip(", ")

    st.session_state.preferences = preferences

    st.caption(f"Active job preferences: {preferences}")


with tabs[3]:
    st.subheader("Cover Letter Generator + Critic Revision Loop")
    st.info(f"Loop halting condition: maximum {MAX_GENERATION_LOOPS} generate-critique-revise passes. Early exit triggers below {int(EARLY_EXIT_SKILL_OVERLAP_THRESHOLD*100)}% skill overlap.")
    
    can_generate = st.session_state.profile is not None and st.session_state.selected_job is not None
    
    if st.button("Generate Application Package", disabled=not can_generate):
        with st.spinner("Generating, critiquing, revising, and sending to judge..."):
            profile_value = st.session_state.profile
            selected_job_value = st.session_state.selected_job
            
            # THE FIX: Removed the strict isinstance() type check
            if profile_value and selected_job_value: 
                try:
                    package = orchestrator.generate_package(profile_value, selected_job_value)
                except ValueError as exc:
                    st.error(str(exc))
                    log_event("Orchestrator", str(exc), "ERROR")
                else:
                    st.session_state.package = package
                    log_event("Orchestrator", f"status={package['status']} loops={package['iteration_count']} judge={package['judge_result']['overall_score']}")
                    st.success(f"Package complete: {package['status']}")
            else:
                st.error("A valid profile and selected job are required before generating a package.")
                log_event("Orchestrator", "Missing profile or job selection", "ERROR")
    if st.session_state.package:
        package = st.session_state.package
        st.markdown("#### Final Cover Letter")
        st.text_area("Generated output", package["final_letter"], height=420)
        st.download_button(
            "Download cover letter TXT",
            data=package["final_letter"],
            file_name="ai_job_application_cover_letter.txt",
            mime="text/plain",
        )
        output_json = json.dumps(package, indent=2)
        st.download_button(
            "Download full agent trace JSON",
            data=output_json,
            file_name="agentic_job_application_trace.json",
            mime="application/json",
        )

with tabs[4]:
    st.subheader("Critic / Reviewer and LLM Judge")
    if not st.session_state.package:
        st.warning("Generate an application package first.")
    else:
        package = st.session_state.package
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### Critic Result")
            critic_df = pd.DataFrame([package["critic_result"]])
            st.dataframe(critic_df, use_container_width=True, hide_index=True)
        with c2:
            st.markdown("#### LLM Judge Result")
            judge_df = pd.DataFrame([package["judge_result"]])
            st.dataframe(judge_df, use_container_width=True, hide_index=True)

        st.markdown("#### Draft History / Revision Trace")
        trace_rows = []
        for d in package["draft_history"]:
            c = d["critic"]
            trace_rows.append({
                "iteration": d["iteration"],
                "status": c.get("status"),
                "professionalism": c.get("professionalism"),
                "alignment": c.get("alignment"),
                "completeness": c.get("completeness"),
                "honesty": c.get("honesty"),
                "harmlessness": c.get("harmlessness"),
                "feedback": c.get("feedback"),
            })
        st.dataframe(pd.DataFrame(trace_rows), use_container_width=True, hide_index=True)

with tabs[5]:
    st.subheader("Evaluation / Risk Dashboard")
    c1, c2 = st.columns([1.2, 1])
    with c1:
        st.markdown("#### Risk Matrix")
        risk_df = pd.DataFrame(RISK_REGISTER)
        st.dataframe(risk_df, use_container_width=True, hide_index=True)
    with c2:
        st.markdown("#### Project Targets")
        target_df = pd.DataFrame([
            {"Metric": "Task Success Rate", "Target": f">= {int(TASK_SUCCESS_TARGET*100)}%", "In-App Measure": "Package status Approved / total runs"},
            {"Metric": "Efficiency", "Target": f"<= {EFFICIENCY_TARGET_AVG_LOOPS} loops", "In-App Measure": "Revision loop count"},
            {"Metric": "Output Quality", "Target": f">= {QUALITY_TARGET}/10", "In-App Measure": "LLM Judge overall score"},
            {"Metric": "HHH", "Target": "Helpful, Honest, Harmless >= 8", "In-App Measure": "Judge H/H/H scores"},
        ])
        st.dataframe(target_df, use_container_width=True, hide_index=True)

    st.markdown("#### Agent Run Log")
    if st.session_state.run_log:
        st.dataframe(pd.DataFrame(st.session_state.run_log), use_container_width=True, hide_index=True)
    else:
        st.caption("No agent events yet.")

    st.markdown("#### Architecture Summary")
    arch_df = pd.DataFrame([
        {"Agent": "Resumé Analyzer", "Input": "PDF/DOCX/TXT resume text", "Output": "Structured profile JSON", "Guardrail": "PII scrub, demographic ignore, injection detection"},
        {"Agent": "Job Matcher", "Input": "Profile + preferences + optional JD", "Output": "Ranked job table", "Guardrail": "Skills-first match; no protected attributes"},
        {"Agent": "Cover Letter Generator", "Input": "Profile + selected job + critic feedback", "Output": "Resume-grounded cover letter", "Guardrail": "Early exit under 25% skill overlap"},
        {"Agent": "Critic / Reviewer", "Input": "Draft + profile + job", "Output": "Scores and revision feedback", "Guardrail": "Max 3 loops; hallucination and PII checks"},
        {"Agent": "LLM Judge", "Input": "Final assets", "Output": "HHH scoring JSON", "Guardrail": "External auditor role separated from orchestrator"},
    ])
    st.dataframe(arch_df, use_container_width=True, hide_index=True)

with tabs[6]:
    st.subheader("Help / Usage Tips")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("""
        #### Quick Start
        1. **Configure API Settings** (Sidebar):
           - Enter your personal **OpenAI API Key** in the sidebar.
           - Toggle **Use Live LLM Mode** ON to enable agent communications with OpenAI models (e.g., `gpt-4o-mini`).
           - *(Optional)* Click **Test Live LLM Connection** to verify your API key works.
        2. **Upload or paste a résumé** in PDF, DOCX, or TXT format.
        3. Optionally **upload or paste a target job description** for exact role alignment.
        4. Click **Extract and Sanitize Inputs** and watch the progress bar complete.
        5. Go to **Resumé Analyzer** and run the analyzer.
        6. Go to **Job Matcher**, select a matched role, then generate the package.
        7. Review **Critic & LLM Judge** scores before using the letter.

        #### Resume Version Control
        - Every extraction creates a new in-session résumé snapshot.
        - Use **Load Selected Version** to compare alternate résumé drafts.
        - Analyzer results are saved back to the active version after analysis.
        - Export the version history JSON if you want a portable audit trail.
        """)

    with c2:
        st.markdown("""
        #### Live LLM Mode vs. Local Fallback
        - **Live LLM Mode (ON):** Sends agent prompts to OpenAI for enhanced resume extraction, tailored cover letters, detailed critique, and objective LLM-judge scoring.
        - **Local Fallback Mode (OFF):** Runs entirely locally using deterministic heuristic logic without sending data to external APIs.

        #### Best Results
        - Use a clean résumé with clear **Skills**, **Projects**, **Experience**, and **Education** sections.
        - Add honest metrics such as ticket volume, project counts, or measurable outcomes.
        - Paste the exact job description when possible; it improves alignment and judge scoring.
        - If the app triggers **low skill overlap**, choose a better-fit job or revise the résumé honestly.

        #### Safety Notes
        - The app strips XML/HTML-like tags and flags prompt-injection signatures.
        - It scrubs common PII from the analyzer context and avoids demographic-based matching.
        - The cover letter generator is instructed not to invent degrees, certifications, metrics, or experience.
        - Final outputs should always be manually reviewed before submission.
        """)

    st.markdown("#### Troubleshooting")
    st.dataframe(pd.DataFrame([
        {"Issue": "LLM mode not enabled or key error", "Fix": "Enter a valid OpenAI API key in the sidebar and ensure 'Use Live LLM Mode' is toggled ON."},
        {"Issue": "Connection test fails", "Fix": "Check that your API key is active, has sufficient OpenAI credits, and has access to the specified model."},
        {"Issue": "PDF text looks incomplete", "Fix": "Use a text-based PDF or paste the resume text manually. Scanned PDFs may need OCR first."},
        {"Issue": "DOC file will not upload", "Fix": "Save legacy .doc files as .docx or PDF before uploading."},
        {"Issue": "Low judge score", "Fix": "Paste a more detailed job description and add truthful role-specific skills or projects to the resume."},
        {"Issue": "Too many revisions", "Fix": "Use the critic feedback to improve the resume/job fit, then create a new resume version."},
    ]), use_container_width=True, hide_index=True)
st.divider()
st.caption("Built as a single-file Streamlit app for the Agentic AI project: multi-agent workflow | front-end | evaluation | risk matrix | pilot-style HHH judging.")
st.divider()
# Fetch the current time and convert to Eastern Time
eastern = pytz.timezone('US/Eastern')
current_time = datetime.now(eastern).strftime("%B %d, %Y - %I:%M %p EST")

# Display it at the bottom using a muted caption style
st.caption(f"⏱️ Dashboard last updated: {current_time}")
